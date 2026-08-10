"""Text extraction for the document formats the LLM node accepts.

Everything funnels through :func:`extract_document`, which returns a
:class:`DocumentPart`: the extracted text plus, for scanned PDF pages that hold no
text at all, rasterised page images to hand to a vision model instead.
"""

from __future__ import annotations

import json
import logging
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field

logger = logging.getLogger("HawkNodes")

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log", ".rst"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx", ".doc"}


@dataclass
class DocumentPart:
    """One parsed document. ``images`` holds data URIs for pages that had no text."""

    name: str
    text: str
    images: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def as_context(self) -> str:
        return f"### {self.name}\n{self.text}".rstrip()


def is_supported(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in SUPPORTED_EXTENSIONS


def parse_page_range(spec: str, total: int) -> list[int]:
    """``"all"`` or a 1-based spec like ``"1-5"`` / ``"1,3,7-9"`` to 0-based indices."""
    spec = (spec or "all").strip().lower()
    if not spec or spec == "all":
        return list(range(total))

    pages: set[int] = set()
    for chunk in spec.split(","):
        chunk = chunk.strip()
        if not chunk:
            continue
        try:
            if "-" in chunk:
                start, end = chunk.split("-", 1)
                for page in range(int(start), int(end) + 1):
                    pages.add(page - 1)
            else:
                pages.add(int(chunk) - 1)
        except ValueError:
            raise ValueError(
                f"Could not read page range {spec!r}. Use 'all', '1-5', or '1,3,7-9'."
            ) from None

    valid = sorted(page for page in pages if 0 <= page < total)
    if not valid:
        raise ValueError(f"Page range {spec!r} selects no pages; the file has {total}.")
    return valid


def _rasterize_pdf_page(path: str, page_index: int) -> str | None:
    """Render one PDF page to a data URI, if PyMuPDF is installed."""
    try:
        import fitz  # PyMuPDF, an optional extra
    except ImportError:
        return None

    try:
        from .images import pil_to_data_uri

        with fitz.open(path) as document:
            pixmap = document[page_index].get_pixmap(dpi=150)
            from PIL import Image
            import io

            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            return pil_to_data_uri(image, "png")
    except Exception as exc:
        logger.warning("HawkNodes: could not rasterize page %d: %s", page_index + 1, exc)
        return None


def _extract_pdf(path: str, pages: str) -> DocumentPart:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError(
            "Reading PDFs needs `pypdf`. Install it with: "
            "pip install -r custom_nodes/HawkNodes/requirements.txt"
        ) from None

    reader = PdfReader(path)
    if reader.is_encrypted:
        try:
            reader.decrypt("")  # many PDFs are "encrypted" with an empty password
        except Exception:
            raise RuntimeError(
                f"{os.path.basename(path)} is password protected; remove the password first."
            ) from None

    indices = parse_page_range(pages, len(reader.pages))
    chunks: list[str] = []
    images: list[str] = []
    warnings: list[str] = []
    empty_pages: list[int] = []

    for index in indices:
        text = (reader.pages[index].extract_text() or "").strip()
        if text:
            chunks.append(f"[page {index + 1}]\n{text}")
            continue

        # No text layer -- most likely a scan. Try to send it as an image instead.
        empty_pages.append(index + 1)
        rendered = _rasterize_pdf_page(path, index)
        if rendered:
            images.append(rendered)

    if empty_pages:
        if images:
            warnings.append(
                f"Pages {_summarize(empty_pages)} have no text layer; sent as images "
                f"instead (needs a vision-capable model)."
            )
        else:
            warnings.append(
                f"Pages {_summarize(empty_pages)} have no text layer and could not be "
                f"rendered. Install PyMuPDF (`pip install pymupdf`) to send scanned "
                f"pages to a vision model."
            )

    return DocumentPart(
        name=os.path.basename(path),
        text="\n\n".join(chunks),
        images=images,
        warnings=warnings,
    )


def _summarize(pages: list[int]) -> str:
    if len(pages) <= 6:
        return ", ".join(str(page) for page in pages)
    return f"{pages[0]}-{pages[-1]} ({len(pages)} pages)"


def _extract_docx(path: str) -> DocumentPart:
    try:
        import docx
    except ImportError:
        raise RuntimeError(
            "Reading .docx needs `python-docx`. Install it with: "
            "pip install -r custom_nodes/HawkNodes/requirements.txt"
        ) from None

    document = docx.Document(path)
    chunks = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables are frequently where the actual content lives in a spec document.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    return DocumentPart(name=os.path.basename(path), text="\n".join(chunks))


def _extract_legacy_doc(path: str) -> DocumentPart:
    """Legacy binary .doc -- python-docx cannot read these, so shell out."""
    name = os.path.basename(path)

    if shutil.which("antiword"):
        result = subprocess.run(
            ["antiword", path], capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and result.stdout.strip():
            return DocumentPart(name=name, text=result.stdout.strip())

    if shutil.which("libreoffice") or shutil.which("soffice"):
        binary = shutil.which("libreoffice") or shutil.which("soffice")
        with tempfile.TemporaryDirectory() as workdir:
            result = subprocess.run(
                [binary, "--headless", "--convert-to", "txt:Text", "--outdir", workdir, path],
                capture_output=True,
                text=True,
                timeout=180,
            )
            converted = os.path.join(workdir, os.path.splitext(name)[0] + ".txt")
            if result.returncode == 0 and os.path.exists(converted):
                with open(converted, "r", encoding="utf-8", errors="replace") as handle:
                    return DocumentPart(name=name, text=handle.read().strip())

    raise RuntimeError(
        f"{name} is a legacy binary .doc, which needs an external converter. "
        f"Either install one (`apt-get install antiword`, or LibreOffice) or "
        f"re-save the file as .docx."
    )


def _extract_plain_text(path: str) -> DocumentPart:
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        text = handle.read()

    if path.lower().endswith(".json"):
        try:
            text = json.dumps(json.loads(text), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            pass  # not valid JSON, send it as-is

    return DocumentPart(name=os.path.basename(path), text=text.strip())


def extract_document(path: str, pages: str = "all") -> DocumentPart:
    """Parse any supported file into a :class:`DocumentPart`."""
    if not os.path.isfile(path):
        raise RuntimeError(f"File not found: {path}")

    extension = os.path.splitext(path)[1].lower()
    if extension == ".pdf":
        part = _extract_pdf(path, pages)
    elif extension == ".docx":
        part = _extract_docx(path)
    elif extension == ".doc":
        part = _extract_legacy_doc(path)
    elif extension in TEXT_EXTENSIONS:
        part = _extract_plain_text(path)
    else:
        raise RuntimeError(
            f"Cannot read {extension or 'a file with no extension'}. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if not part.text and not part.images:
        part.warnings.append(f"{part.name} produced no text at all.")
    return part


def assemble_context(
    parts: list[DocumentPart], extra_text: str = "", max_chars: int = 0
) -> tuple[str, list[str]]:
    """Join documents into one context block, truncating to ``max_chars``.

    Returns the block and any warnings, so the node can surface them on its
    ``context_used`` output instead of failing silently.
    """
    warnings: list[str] = []
    blocks: list[str] = []

    if extra_text and extra_text.strip():
        blocks.append(extra_text.strip())
    for part in parts:
        warnings.extend(part.warnings)
        if part.text:
            blocks.append(part.as_context())

    context = "\n\n".join(blocks)
    if max_chars > 0 and len(context) > max_chars:
        dropped = len(context) - max_chars
        context = context[:max_chars] + f"\n\n...[truncated {dropped} chars]..."
        warnings.append(
            f"Context hit the {max_chars} char limit; {dropped} chars were dropped. "
            f"Raise `context_chars` to send more."
        )
    return context, warnings
