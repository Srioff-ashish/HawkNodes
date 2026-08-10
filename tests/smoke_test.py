#!/usr/bin/env python3
"""Exercise the Atlas client without ComfyUI.

Nothing here imports ComfyUI, so a failure points at the Atlas integration rather
than at the node wiring.

    python tests/smoke_test.py --key $ATLAS_API_KEY            # offline + chat
    python tests/smoke_test.py --offline                       # no network at all
    python tests/smoke_test.py --key ... --t2i --i2i           # image generation too

Every step is optional and reports pass/fail independently, so one broken model id
does not hide the rest.
"""

from __future__ import annotations

import argparse
import asyncio
import base64
import io
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from hawknodes import client, documents, images  # noqa: E402

PASS, FAIL, SKIP = "  ok  ", " FAIL ", " skip "
results: list[tuple[str, str, str]] = []


def record(name: str, status: str, detail: str = "") -> None:
    results.append((name, status, detail))
    print(f"[{status}] {name}" + (f" -- {detail}" if detail else ""), flush=True)


def sample_image_bytes(size=(64, 48), color=(220, 40, 40)) -> bytes:
    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", size, color).save(buffer, format="PNG")
    return buffer.getvalue()


# ------------------------------------------------------------------ offline


def test_url_normalization() -> None:
    cases = [
        ("", client.DEFAULT_CHAT_URL),
        ("https://api.atlascloud.ai", "https://api.atlascloud.ai/v1"),
        ("https://api.atlascloud.ai/v1", "https://api.atlascloud.ai/v1"),
        ("https://api.atlascloud.ai/v1/chat/completions", "https://api.atlascloud.ai/v1"),
    ]
    for given, expected in cases:
        actual = client.normalize_chat_url(given)
        assert actual == expected, f"{given!r} -> {actual!r}, expected {expected!r}"

    for given in ("https://api.atlascloud.ai", "https://api.atlascloud.ai/api/v1",
                  "https://api.atlascloud.ai/api", "https://api.atlascloud.ai/v1"):
        actual = client.normalize_image_base(given)
        assert actual == "https://api.atlascloud.ai", f"{given!r} -> {actual!r}"
    record("url normalization", PASS)


def test_key_resolution() -> None:
    saved = os.environ.pop("ATLAS_API_KEY", None)
    try:
        assert client.resolve_api_key("  widget-key  ") == "widget-key"
        os.environ["ATLAS_API_KEY"] = "env-key"
        assert client.resolve_api_key("") == "env-key"
        assert client.resolve_api_key("widget-key") == "widget-key"
        os.environ.pop("ATLAS_API_KEY")
        try:
            client.resolve_api_key("")
        except client.AtlasError as exc:
            assert "ATLAS_API_KEY" in str(exc)
        else:
            raise AssertionError("a missing key should raise")
    finally:
        if saved is not None:
            os.environ["ATLAS_API_KEY"] = saved
    record("api key resolution", PASS)


def test_page_ranges() -> None:
    assert documents.parse_page_range("all", 5) == [0, 1, 2, 3, 4]
    assert documents.parse_page_range("1-3", 10) == [0, 1, 2]
    assert documents.parse_page_range("1,3,5-6", 10) == [0, 2, 4, 5]
    assert documents.parse_page_range("8-99", 10) == [7, 8, 9]  # clamped
    for bad in ("abc", "1-x"):
        try:
            documents.parse_page_range(bad, 10)
        except ValueError:
            pass
        else:
            raise AssertionError(f"{bad!r} should be rejected")
    record("pdf page ranges", PASS)


def test_text_extraction() -> None:
    with tempfile.TemporaryDirectory() as workdir:
        path = os.path.join(workdir, "notes.txt")
        with open(path, "w", encoding="utf-8") as handle:
            handle.write("hello from a text file")
        part = documents.extract_document(path)
        assert part.text == "hello from a text file", part.text
        assert part.name == "notes.txt"

        json_path = os.path.join(workdir, "data.json")
        with open(json_path, "w", encoding="utf-8") as handle:
            handle.write('{"b":2,"a":1}')
        assert '"a": 1' in documents.extract_document(json_path).text

        bad_path = os.path.join(workdir, "clip.mp4")
        open(bad_path, "wb").close()
        try:
            documents.extract_document(bad_path)
        except RuntimeError as exc:
            assert ".pdf" in str(exc)
        else:
            raise AssertionError("an unsupported extension should raise")
    record("text/json extraction", PASS)


def test_docx_extraction() -> None:
    try:
        import docx
    except ImportError:
        record("docx extraction", SKIP, "python-docx not installed")
        return

    with tempfile.TemporaryDirectory() as workdir:
        path = os.path.join(workdir, "spec.docx")
        document = docx.Document()
        document.add_paragraph("First paragraph.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Key"
        table.rows[0].cells[1].text = "Value"
        document.save(path)

        part = documents.extract_document(path)
        assert "First paragraph." in part.text
        assert "Key | Value" in part.text, part.text
    record("docx extraction", PASS, f"{len(part.text)} chars")


def test_pdf_extraction() -> None:
    try:
        from pypdf import PdfWriter
    except ImportError:
        record("pdf extraction", SKIP, "pypdf not installed")
        return

    with tempfile.TemporaryDirectory() as workdir:
        path = os.path.join(workdir, "blank.pdf")
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        writer.add_blank_page(width=200, height=200)
        with open(path, "wb") as handle:
            writer.write(handle)

        # Blank pages have no text layer, which is exactly the scanned-PDF path:
        # it must warn rather than silently return nothing.
        part = documents.extract_document(path, "all")
        assert part.warnings, "a text-free PDF should produce a warning"
        assert documents.parse_page_range("2", 2) == [1]
    record("pdf extraction", PASS, part.warnings[0][:60])


def test_context_assembly() -> None:
    parts = [
        documents.DocumentPart(name="a.txt", text="A" * 100),
        documents.DocumentPart(name="b.txt", text="B" * 100),
    ]
    context, warnings = documents.assemble_context(parts, "prefix", max_chars=0)
    assert context.startswith("prefix")
    assert "### a.txt" in context and "### b.txt" in context
    assert not warnings

    context, warnings = documents.assemble_context(parts, "", max_chars=50)
    assert "truncated" in context
    assert warnings and "context_chars" in warnings[0]
    record("context assembly + truncation", PASS)


def test_image_roundtrip() -> None:
    try:
        import torch
    except ImportError:
        record("image round trip", SKIP, "torch not installed")
        return

    tensor = torch.rand(2, 128, 256, 3)  # a batch of two
    uris = images.tensor_batch_to_data_uris(tensor, max_side=64, image_format="png")
    assert len(uris) == 2
    assert uris[0].startswith("data:image/png;base64,")

    from PIL import Image

    decoded = Image.open(io.BytesIO(base64.b64decode(uris[0].split(",", 1)[1])))
    assert max(decoded.size) == 64, decoded.size  # downscaled

    back = images.stack_images([sample_image_bytes(), sample_image_bytes((32, 32))])
    assert back.shape[0] == 2, back.shape
    assert back.shape[1:3] == (48, 64), back.shape  # second image resized to match
    assert 0.0 <= float(back.min()) and float(back.max()) <= 1.0
    record("image round trip", PASS, f"batch {tuple(back.shape)}")


def test_output_decoding() -> None:
    raw = sample_image_bytes()
    encoded = base64.b64encode(raw).decode()
    assert client._decode_output(f"data:image/png;base64,{encoded}") == raw
    assert client._decode_output(encoded) == raw
    assert client._decode_output("https://example.com/a.png") == "https://example.com/a.png"
    record("output decoding (data uri / base64 / url)", PASS)


def test_message_shape() -> None:
    from hawknodes.nodes import llm  # imports comfy_api; only works inside ComfyUI

    messages = llm._build_messages("sys", "hello", [], "auto")
    assert messages[0]["role"] == "system"
    assert isinstance(messages[1]["content"], str), "text-only must send a plain string"

    messages = llm._build_messages("", "hello", ["data:image/png;base64,AA"], "high")
    assert isinstance(messages[0]["content"], list)
    assert messages[0]["content"][1]["image_url"]["detail"] == "high"
    record("message shape", PASS)


def test_image_payload() -> None:
    from hawknodes import registry
    from hawknodes.nodes import image_base

    # gpt-image-2: every widget it declares is sent.
    payload = image_base.build_payload(
        {"model": "openai/gpt-image-2/edit", "size": "1024x1536",
         "quality": "high", "input_fidelity": "low"},
        "a hawk",
        seed=7, negative_prompt="blurry", n=2, output_format="png",
        extra_params='{"guidance_scale": 3.5}',
        image_uris=["data:image/png;base64,AA"],
    )
    assert payload["enable_base64_output"] is True
    assert payload["enable_sync_mode"] is False
    for key, value in [("size", "1024x1536"), ("quality", "high"),
                       ("input_fidelity", "low"), ("seed", 7),
                       ("negative_prompt", "blurry"), ("n", 2),
                       ("guidance_scale", 3.5)]:
        assert payload.get(key) == value, f"{key}={payload.get(key)!r}"
    assert payload["images"] == ["data:image/png;base64,AA"]

    # A generic model on defaults must send nothing it was not asked to.
    payload = image_base.build_payload(
        {"model": "seedream-3.0", "size": registry.SIZE_DEFAULT},
        "a hawk",
        seed=0, negative_prompt="", n=1, output_format="jpeg", extra_params="",
    )
    for absent in ("size", "quality", "input_fidelity", "seed", "negative_prompt", "n", "images"):
        assert absent not in payload, f"{absent} should have been omitted, got {payload}"

    # model_override wins; `custom` without one is a clear error.
    assert registry.resolve_slug({"model": "a", "model_override": " b "}) == "b"
    assert registry.resolve_slug({"model": "a"}) == "a"
    try:
        registry.resolve_slug({"model": registry.CUSTOM_SLUG, "model_override": ""})
    except ValueError as exc:
        assert "model_override" in str(exc)
    else:
        raise AssertionError("custom with no override should raise")

    for bad, needle in [("not json", "not valid JSON"), ("[1,2]", "JSON object")]:
        try:
            image_base.build_payload(
                {"model": "m"}, "p", seed=0, negative_prompt="", n=1,
                output_format="jpeg", extra_params=bad,
            )
        except ValueError as exc:
            assert needle in str(exc), str(exc)
        else:
            raise AssertionError(f"{bad!r} should be rejected")

    try:
        image_base.build_payload({"model": "m"}, "   ", seed=0, negative_prompt="",
                                 n=1, output_format="jpeg", extra_params="")
    except ValueError:
        pass
    else:
        raise AssertionError("an empty prompt should be rejected")
    record("image payload construction", PASS)


# ------------------------------------------------------------------- network


async def test_list_models(api_url: str, key: str) -> list[str]:
    models = await client.fetch_models(api_url, key)
    if models:
        record("GET /v1/models", PASS, f"{len(models)} models, e.g. {models[0]}")
    else:
        record(
            "GET /v1/models",
            SKIP,
            "no models returned; the dropdown falls back to models.json",
        )
    return models


async def test_chat(api_url: str, key: str, model: str) -> None:
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": "Reply with exactly: HAWK OK"}],
        "max_tokens": 20,
        "temperature": 0.0,
        "stream": False,
    }
    try:
        response = await client.chat_completion(api_url, key, payload, timeout=60)
        text = client.extract_message_text(response)
        record("chat completion", PASS, f"{model} -> {text.strip()[:40]!r}")
    except Exception as exc:
        record("chat completion", FAIL, f"{model}: {exc}")


async def test_vision(api_url: str, key: str, model: str) -> None:
    encoded = base64.b64encode(sample_image_bytes()).decode()
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "What colour is this image? One word."},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{encoded}"},
                    },
                ],
            }
        ],
        "max_tokens": 20,
        "stream": False,
    }
    try:
        response = await client.chat_completion(api_url, key, payload, timeout=90)
        text = client.extract_message_text(response)
        record("vision chat", PASS, f"{model} -> {text.strip()[:40]!r}")
    except Exception as exc:
        record("vision chat", FAIL, f"{model}: {exc}")


async def test_t2i(api_url: str, key: str, model: str, outdir: str) -> str | None:
    payload = {
        "model": model,
        "prompt": "a red hawk silhouette on a white background, minimal vector logo",
        "output_format": "png",
        "enable_base64_output": True,
        "enable_sync_mode": False,
    }
    try:
        results_ = await client.generate_image(api_url, key, payload, timeout=300)
        path = os.path.join(outdir, "t2i.png")
        with open(path, "wb") as handle:
            handle.write(results_[0])
        record("text to image", PASS, f"{model} -> {path} ({len(results_[0])} bytes)")
        return path
    except Exception as exc:
        record("text to image", FAIL, f"{model}: {exc}")
        return None


async def test_i2i(api_url: str, key: str, model: str, source: str, outdir: str) -> None:
    with open(source, "rb") as handle:
        encoded = base64.b64encode(handle.read()).decode()
    payload = {
        "model": model,
        "prompt": "make the background deep blue",
        "images": [f"data:image/png;base64,{encoded}"],
        "output_format": "png",
        "enable_base64_output": True,
        "enable_sync_mode": False,
    }
    try:
        results_ = await client.generate_image(api_url, key, payload, timeout=300)
        path = os.path.join(outdir, "i2i.png")
        with open(path, "wb") as handle:
            handle.write(results_[0])
        record("image to image", PASS, f"{model} -> {path}")
    except Exception as exc:
        record("image to image", FAIL, f"{model}: {exc}")


# ---------------------------------------------------------------------- main


async def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--key", default=os.environ.get("ATLAS_API_KEY", ""))
    parser.add_argument("--api-url", default=client.DEFAULT_CHAT_URL)
    parser.add_argument("--image-url", default=client.DEFAULT_IMAGE_URL)
    parser.add_argument("--model", default="deepseek-v3")
    parser.add_argument("--vision-model", default="", help="a vision-capable model id")
    parser.add_argument("--t2i-model", default="openai/gpt-image-2/text-to-image")
    parser.add_argument("--i2i-model", default="openai/gpt-image-2/edit")
    parser.add_argument("--offline", action="store_true", help="skip all network calls")
    parser.add_argument("--t2i", action="store_true", help="run a real generation")
    parser.add_argument("--i2i", action="store_true", help="run a real edit")
    parser.add_argument("--outdir", default=tempfile.gettempdir())
    args = parser.parse_args()

    print("--- offline ---")
    for check in (
        test_url_normalization,
        test_key_resolution,
        test_page_ranges,
        test_text_extraction,
        test_docx_extraction,
        test_pdf_extraction,
        test_context_assembly,
        test_image_roundtrip,
        test_output_decoding,
    ):
        try:
            check()
        except Exception as exc:
            record(check.__name__, FAIL, str(exc))

    for check in (test_message_shape, test_image_payload):
        try:
            check()
        except ImportError:
            record(check.__name__, SKIP, "needs ComfyUI on the path")
        except Exception as exc:
            record(check.__name__, FAIL, str(exc))

    if not args.offline:
        if not args.key:
            print("\nNo API key given; skipping network checks. Pass --key or set ATLAS_API_KEY.")
        else:
            print("\n--- network ---")
            await test_list_models(args.api_url, args.key)
            await test_chat(args.api_url, args.key, args.model)
            if args.vision_model:
                await test_vision(args.api_url, args.key, args.vision_model)
            if args.t2i or args.i2i:
                generated = await test_t2i(args.image_url, args.key, args.t2i_model, args.outdir)
                if args.i2i and generated:
                    await test_i2i(args.image_url, args.key, args.i2i_model, generated, args.outdir)

    failed = [name for name, status, _ in results if status == FAIL]
    print(f"\n{len(results) - len(failed)}/{len(results)} checks passed.")
    if failed:
        print("Failed: " + ", ".join(failed))
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
